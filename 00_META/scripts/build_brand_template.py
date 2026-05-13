#!/usr/bin/env python3
"""
Build the Mejia, IA & CIA brand DOCX template.

Meets the first-meeting obligation:
  - Header with client company name field + firm wordmark
  - Footer with page number + firm name
  - Brand fonts (Manrope preferred, Calibri fallback)
  - Brand colors (navy + gold + paper background tones)
  - Demonstrates a real Skill 1 (Analisis Compania) report shape

Output: 2026-05-11_brand-template-Analisis_Compania.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

# --- brand tokens ---
NAVY   = RGBColor(0x0a, 0x1f, 0x33)
GOLD   = RGBColor(0xd4, 0xb1, 0x6a)
INK    = RGBColor(0x1c, 0x25, 0x30)
MUTED  = RGBColor(0x6b, 0x75, 0x85)
PAPER  = "FAF9F5"   # background hex (no leading #)
LINE   = "E8E3D3"
FONT_HEADING = "Manrope"   # falls back to default if not installed
FONT_BODY    = "Manrope"
FONT_FALLBACK = "Calibri"

OUTPUT = "/home/ricardo/brain/01_VENTURES/04_MejiaIACia/00_META/proposals/2026-05-11_brand-template-Analisis_Compania.docx"


# --- helpers ---
def set_run(run, *, font=FONT_BODY, size=11, bold=False, color=INK, spacing_pts=None, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # also set East Asian font so Word/LO doesn't fall back
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:cs'), font)
    if spacing_pts is not None:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:val'), str(int(spacing_pts * 20)))  # twips
        rpr.append(spacing)
    if caps:
        c = OxmlElement('w:caps')
        c.set(qn('w:val'), '1')
        rpr.append(c)


def add_field(paragraph, instr_text):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = instr_text
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)
    return run


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_bottom_border(paragraph, color_hex="d4b16a", size_pt=2):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(size_pt * 8)))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=6, line=1.4):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line


# --- build the document ---
doc = Document()

# A4 margins
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

# different first page (no header/footer on cover)
section = doc.sections[0]
section.different_first_page_header_footer = True

# === HEADER (recurring) ===
hdr = section.header
hdr_p = hdr.paragraphs[0]
hdr_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.6), WD_TAB_ALIGNMENT.RIGHT)
# left: M.IA.C wordmark
run_l = hdr_p.add_run("M")
set_run(run_l, font=FONT_HEADING, size=12, bold=True, color=NAVY)
run_dot1 = hdr_p.add_run("·")
set_run(run_dot1, font=FONT_HEADING, size=12, bold=True, color=GOLD)
run_ia = hdr_p.add_run("IA")
set_run(run_ia, font=FONT_HEADING, size=12, bold=True, color=NAVY)
run_dot2 = hdr_p.add_run("·")
set_run(run_dot2, font=FONT_HEADING, size=12, bold=True, color=GOLD)
run_c = hdr_p.add_run("C")
set_run(run_c, font=FONT_HEADING, size=12, bold=True, color=NAVY)
# right: client name + report type
hdr_p.add_run("\t")
run_r = hdr_p.add_run("Análisis Compañía  ·  ")
set_run(run_r, font=FONT_BODY, size=9, color=MUTED, caps=True, spacing_pts=1.5)
run_client = hdr_p.add_run("[ NOMBRE DEL CLIENTE ]")
set_run(run_client, font=FONT_BODY, size=9, bold=True, color=NAVY, caps=True, spacing_pts=1.5)
add_bottom_border(hdr_p, color_hex="e8e3d3", size_pt=0.75)

# === FOOTER (recurring) ===
ftr = section.footer
ftr_p = ftr.paragraphs[0]
ftr_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.6), WD_TAB_ALIGNMENT.RIGHT)
# left: firm + tag
run_f1 = ftr_p.add_run("Mejía, IA & CIA")
set_run(run_f1, font=FONT_BODY, size=9, bold=True, color=NAVY)
run_f2 = ftr_p.add_run("   ·   Reporte confidencial")
set_run(run_f2, font=FONT_BODY, size=9, color=MUTED)
# right: page X de Y
ftr_p.add_run("\t")
run_p1 = ftr_p.add_run("Página ")
set_run(run_p1, font=FONT_BODY, size=9, color=MUTED)
run_p2 = add_field(ftr_p, "PAGE")
set_run(run_p2, font=FONT_BODY, size=9, bold=True, color=NAVY)
run_p3 = ftr_p.add_run(" de ")
set_run(run_p3, font=FONT_BODY, size=9, color=MUTED)
run_p4 = add_field(ftr_p, "NUMPAGES")
set_run(run_p4, font=FONT_BODY, size=9, bold=True, color=NAVY)

# === COVER PAGE (first page header/footer are empty by default since different_first_page is on) ===

# big navy block at top of cover
cover_block = doc.add_paragraph()
shade_paragraph(cover_block, "0A1F33")
cover_block.paragraph_format.space_before = Pt(0)
cover_block.paragraph_format.space_after = Pt(0)
r = cover_block.add_run("    M · IA · C    ")
set_run(r, font=FONT_HEADING, size=32, bold=True, color=GOLD)
# vertical breathing room
for _ in range(3):
    cover_block.add_run("\n")

# spacer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(60)
spacer.paragraph_format.space_after = Pt(0)

# eyebrow
eyebrow = doc.add_paragraph()
set_paragraph_spacing(eyebrow, before_pt=0, after_pt=8)
r = eyebrow.add_run("REPORTE  ·  SKILL 1")
set_run(r, font=FONT_BODY, size=10, bold=True, color=GOLD, caps=True, spacing_pts=4)

# title
title = doc.add_paragraph()
set_paragraph_spacing(title, before_pt=0, after_pt=12, line=1.05)
r = title.add_run("Análisis Compañía")
set_run(r, font=FONT_HEADING, size=44, bold=True, color=NAVY)

# subtitle / client placeholder
sub = doc.add_paragraph()
set_paragraph_spacing(sub, before_pt=0, after_pt=24, line=1.2)
r = sub.add_run("[ Nombre del cliente ]")
set_run(r, font=FONT_HEADING, size=22, bold=False, color=MUTED)

# rule
rule = doc.add_paragraph()
add_bottom_border(rule, color_hex="d4b16a", size_pt=1.5)
set_paragraph_spacing(rule, before_pt=0, after_pt=20)

# preparado por
prep = doc.add_paragraph()
set_paragraph_spacing(prep, before_pt=0, after_pt=4)
r = prep.add_run("Preparado por")
set_run(r, font=FONT_BODY, size=10, color=MUTED, caps=True, spacing_pts=2)
firm = doc.add_paragraph()
set_paragraph_spacing(firm, before_pt=0, after_pt=18)
r = firm.add_run("Mejía, IA & CIA")
set_run(r, font=FONT_BODY, size=13, bold=True, color=NAVY)

# date
date_p = doc.add_paragraph()
set_paragraph_spacing(date_p, before_pt=0, after_pt=4)
r = date_p.add_run("Fecha")
set_run(r, font=FONT_BODY, size=10, color=MUTED, caps=True, spacing_pts=2)
date_v = doc.add_paragraph()
set_paragraph_spacing(date_v, before_pt=0, after_pt=18)
r = date_v.add_run("[ DD / MM / AAAA ]")
set_run(r, font=FONT_BODY, size=13, color=INK)

# version
ver_p = doc.add_paragraph()
set_paragraph_spacing(ver_p, before_pt=0, after_pt=4)
r = ver_p.add_run("Versión")
set_run(r, font=FONT_BODY, size=10, color=MUTED, caps=True, spacing_pts=2)
ver_v = doc.add_paragraph()
set_paragraph_spacing(ver_v, before_pt=0, after_pt=0)
r = ver_v.add_run("v1.0")
set_run(r, font=FONT_BODY, size=13, color=INK)

# Page break to body
doc.add_page_break()

# === BODY ===

def h1(text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=18, after_pt=8, line=1.1)
    r = p.add_run(text)
    set_run(r, font=FONT_HEADING, size=22, bold=True, color=NAVY)
    add_bottom_border(p, color_hex="d4b16a", size_pt=1)
    return p

def h2(text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=14, after_pt=4)
    r = p.add_run(text)
    set_run(r, font=FONT_HEADING, size=14, bold=True, color=NAVY)
    return p

def body(text, *, color=INK):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=8, line=1.45)
    r = p.add_run(text)
    set_run(r, font=FONT_BODY, size=11, color=color)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before_pt=0, after_pt=4, line=1.4)
    for r in p.runs:
        set_run(r, font=FONT_BODY, size=11, color=INK)
    r = p.add_run(text)
    set_run(r, font=FONT_BODY, size=11, color=INK)
    return p

def callout(label, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "FAF4E1")
    set_paragraph_spacing(p, before_pt=10, after_pt=10, line=1.4)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4); pf.right_indent = Cm(0.4)
    r = p.add_run(label.upper() + "  ")
    set_run(r, font=FONT_BODY, size=9, bold=True, color=GOLD, caps=True, spacing_pts=2)
    r = p.add_run(text)
    set_run(r, font=FONT_BODY, size=11, color=INK)
    return p

# Resumen ejecutivo
h1("Resumen ejecutivo")
body("Este reporte sintetiza los hallazgos del Skill 1 — Análisis Compañía — aplicado a [Nombre del cliente]. Cubre identidad corporativa, posicionamiento de mercado, madurez digital observable y señales de procesos manuales susceptibles de automatización con IA.")
callout("Conclusión preliminar",
        "[Una a tres oraciones que resumen la oportunidad principal detectada. Por ejemplo: la empresa opera un proceso de atención al cliente fragmentado entre tres números de WhatsApp y un formulario web sin respuesta automática — Quick Win prioritario.]")

# Identidad
h1("1. Identidad corporativa")
h2("Datos básicos")
body("[ Razón social ] · [ NIT ] · [ Cámara de Comercio ] · [ Fecha de constitución ] · [ Tipo societario ].")
h2("Productos y servicios")
body("[ Descripción de la oferta principal según el sitio web oficial. Si la empresa tiene múltiples líneas de negocio, listar las tres principales y la proporción aproximada de cada una. ]")
h2("Presencia digital")
bullet("Sitio web: [URL] — calidad visual [profesional / amateur / desactualizado].")
bullet("LinkedIn: [N empleados] · [última actividad].")
bullet("Google Reviews: [calificación X.X/5] · [N reseñas].")

# Madurez digital
h1("2. Madurez digital observable")
body("Escala 1–5 según el marco de la firma. La calificación se basa en señales públicas observables (sitio web, redes sociales, presencia en plataformas) y no requiere acceso interno.")
callout("Calificación", "[ X / 5 ] — [breve justificación de la calificación]")

# Quick Wins
h1("3. Señales de procesos manuales — Quick Wins")
body("Los siguientes patrones son señales externas de procesos que típicamente se automatizan con IA en un plazo corto:")
bullet("[ Señal observada 1 — ej: tres números de WhatsApp distintos en el sitio web ].")
bullet("[ Señal observada 2 — ej: formulario de contacto sin respuesta automática ].")
bullet("[ Señal observada 3 — ej: catálogo descargable en PDF, sin e-commerce ].")
callout("Quick Win prioritario",
        "[ Descripción específica para esta empresa, no genérica del sector. Incluir: qué se automatiza, con qué herramienta, esfuerzo estimado y resultado esperado. ]")

# Próximos pasos
h1("4. Próximos pasos")
body("La progresión natural desde el Skill 1 hacia el resto de la cadena:")
bullet("Skill 2 — Análisis Industria: contexto sectorial y benchmarks.")
bullet("Skill 3 — Análisis Competencia: posicionamiento relativo.")
bullet("Skill 4 — Guión Reunión: preparación para la primera reunión.")
bullet("Skill 5 — Reporte Oportunidades: entregable de cierre con propuesta.")

# Anexo / fuentes
h1("5. Fuentes consultadas")
body("Todo dato citado en este reporte proviene de las siguientes fuentes públicas:")
bullet("Sitio web oficial — [URL] — consultado el [fecha].")
bullet("LinkedIn de la empresa — [URL] — consultado el [fecha].")
bullet("Google Reviews — [N reseñas analizadas] — consultado el [fecha].")
bullet("RUES / Cámara de Comercio — [si aplica].")
bullet("Prensa especializada — [Portafolio / Dinero / La República si aplica].")

# Save
doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
