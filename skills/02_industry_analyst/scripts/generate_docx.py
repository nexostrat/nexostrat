#!/usr/bin/env python3
"""
generate_docx.py — Converts an industry-analyst .md report to a styled .docx

Usage:
    python generate_docx.py <input.md> <output.docx>

Requires: python-docx
    pip install python-docx --break-system-packages
"""

import sys
import re
from pathlib import Path
from datetime import datetime

# Brand surface (palette, logos, header/footer helpers)
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "shared"))
try:
    import brand
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


def parse_md_to_blocks(md_text):
    """Parse markdown into structured blocks for docx rendering."""
    blocks = []
    lines = md_text.split('\n')
    i = 0

    # Detect and skip YAML frontmatter
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1

    # Track whether we're inside a table
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows (---|---) 
            if re.match(r'^\|[-| :]+\|$', line.strip()):
                i += 1
                continue
            # Parse table row
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                blocks.append({'type': 'table', 'rows': table_rows})
                table_rows = []
                in_table = False

        stripped = line.strip()

        if not stripped:
            blocks.append({'type': 'empty'})
        elif stripped.startswith('# ') and not stripped.startswith('##'):
            blocks.append({'type': 'h1', 'text': stripped[2:].strip()})
        elif stripped.startswith('## '):
            blocks.append({'type': 'h2', 'text': stripped[3:].strip()})
        elif stripped.startswith('### '):
            blocks.append({'type': 'h3', 'text': stripped[4:].strip()})
        elif stripped.startswith('#### '):
            blocks.append({'type': 'h4', 'text': stripped[5:].strip()})
        elif stripped.startswith('- ') or stripped.startswith('* '):
            blocks.append({'type': 'bullet', 'text': stripped[2:].strip()})
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            blocks.append({'type': 'numbered', 'text': text})
        elif stripped.startswith('```'):
            # Skip code fences
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            blocks.append({'type': 'bold_line', 'text': stripped.strip('*')})
        elif stripped.startswith('---'):
            blocks.append({'type': 'hr'})
        else:
            blocks.append({'type': 'paragraph', 'text': stripped})

        i += 1

    if in_table and table_rows:
        blocks.append({'type': 'table', 'rows': table_rows})

    return blocks


def clean_markdown_inline(text):
    """Remove inline markdown formatting for plain text in docx."""
    # Bold+italic
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    # Bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Inline code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Links [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text


def add_run_with_formatting(paragraph, text):
    """Add a run to a paragraph, handling **bold** inline markers."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            clean = re.sub(r'\*(.*?)\*', r'\1', part)
            clean = re.sub(r'`(.*?)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean)
            if clean:
                paragraph.add_run(clean)


def generate_docx(md_path, output_path):
    """Generate a styled .docx from an industry analyst markdown report."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
        sys.exit(1)

    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_md_to_blocks(md_text)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Local aliases — pinned to brand module (single source of truth)
    DARK_BLUE = brand.MIDNIGHT_BLUE
    ACCENT    = brand.SKY_BLUE
    MID_GRAY  = brand.GRAY_500
    LIGHT_BG  = RGBColor(0xF4, 0xF7, 0xFB)  # Local: alt table-row tint
    BLACK     = brand.BLACK

    # ── Style helpers ─────────────────────────────────────────────────────────
    def set_para_spacing(para, before=0, after=4):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    def add_h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        set_para_spacing(p, before=12, after=8)
        return p

    def add_h2(text):
        # Strip section number prefix for cleaner rendering
        clean = clean_markdown_inline(text)
        p = doc.add_paragraph()
        run = p.add_run(clean)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        set_para_spacing(p, before=14, after=3)
        # Bottom border via XML
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), brand.HEX_SKY_BLUE)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = MID_GRAY
        set_para_spacing(p, before=8, after=2)
        return p

    def add_h4(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BLACK
        set_para_spacing(p, before=4, after=2)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=1, after=3)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=0, after=2)
        return p

    def add_numbered(text):
        p = doc.add_paragraph(style='List Number')
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=0, after=2)
        return p

    def shade_cell(cell, hex_color='F4F7FB'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_table_block(rows):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= num_cols:
                    break
                cell = row.cells[c_idx]
                clean = clean_markdown_inline(cell_text)
                p = cell.paragraphs[0]
                run = p.add_run(clean)
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shade_cell(cell, brand.HEX_MIDNIGHT_BLUE)
                else:
                    run.font.color.rgb = BLACK
                    if r_idx % 2 == 0:
                        shade_cell(cell, 'EBF3FA')
                    else:
                        shade_cell(cell, 'FFFFFF')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()  # spacing after table

    # ── Title page ────────────────────────────────────────────────────────────
    # Extract title from first H1 in blocks
    title_text = "Análisis de Industria"
    subtitle_text = "Colombia"
    for block in blocks[:5]:
        if block['type'] == 'h1':
            title_text = clean_markdown_inline(block['text'])
            break

    # Look for subtitle (the bold date line)
    for block in blocks[:8]:
        if block.get('text', '').startswith('Colombia') or block.get('text', '').startswith('**Colombia'):
            subtitle_text = clean_markdown_inline(block.get('text', ''))
            break

    # Brand logo at top of cover (via brand module)
    brand.apply_cover_logo(doc, width_inches=3.8, space_after_pt=36)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title_text)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(subtitle_text)
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = MID_GRAY
    p_sub.paragraph_format.space_after = Pt(4)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"Nexostrat — Uso Interno  ·  {datetime.now().strftime('%d %b %Y')}")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = MID_GRAY

    doc.add_page_break()

    # ── Body ─────────────────────────────────────────────────────────────────
    skip_title = True  # Skip the H1 we already used for title page
    for block in blocks:
        btype = block['type']

        if btype == 'h1':
            if skip_title:
                skip_title = False
                continue
            add_h1(block['text'])
        elif btype == 'h2':
            add_h2(block['text'])
        elif btype == 'h3':
            add_h3(block['text'])
        elif btype == 'h4':
            add_h4(block['text'])
        elif btype == 'bullet':
            add_bullet(block['text'])
        elif btype == 'numbered':
            add_numbered(block['text'])
        elif btype == 'table':
            add_table_block(block['rows'])
        elif btype == 'bold_line':
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(block['text']))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
            set_para_spacing(p, before=4, after=2)
        elif btype == 'hr':
            p = doc.add_paragraph()
            set_para_spacing(p, before=2, after=2)
        elif btype == 'paragraph':
            if block['text']:
                add_body(block['text'])
        elif btype == 'empty':
            pass  # Don't add blank paragraphs for every empty line

    # Brand header + footer (via brand module — skips cover automatically)
    brand.apply_brand_header(doc)
    brand.apply_brand_footer(doc)

    doc.save(output_path)
    print(f"✅ DOCX generado: {output_path}")
    size_kb = Path(output_path).stat().st_size // 1024
    print(f"   Tamaño: {size_kb} KB")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Uso: python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = sys.argv[1]
    out_path = sys.argv[2]

    if not Path(md_path).exists():
        print(f"ERROR: No se encontró el archivo: {md_path}")
        sys.exit(1)

    generate_docx(md_path, out_path)
