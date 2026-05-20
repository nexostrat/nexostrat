#!/usr/bin/env python3
"""
generate_docx.py — Converts a competitor-analyst .md report to a styled .docx

Usage:
    python generate_docx.py <input.md> <output.docx>

Requires: python-docx
    pip install python-docx --break-system-packages
"""

import sys
import re
from pathlib import Path
from datetime import datetime

# Brand surface (palette, logos, header/footer helpers)
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "shared"))
try:
    import brand
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


def parse_md_to_blocks(md_text):
    blocks = []
    lines = md_text.split('\n')
    i = 0
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1
    in_table = False
    table_rows = []
    while i < len(lines):
        line = lines[i]
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^\|[-| :]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                blocks.append({'type': 'table', 'rows': table_rows})
                table_rows = []
                in_table = False
        stripped = line.strip()
        if not stripped:
            blocks.append({'type': 'empty'})
        elif stripped.startswith('# ') and not stripped.startswith('##'):
            blocks.append({'type': 'h1', 'text': stripped[2:].strip()})
        elif stripped.startswith('## '):
            blocks.append({'type': 'h2', 'text': stripped[3:].strip()})
        elif stripped.startswith('### '):
            blocks.append({'type': 'h3', 'text': stripped[4:].strip()})
        elif stripped.startswith('#### '):
            blocks.append({'type': 'h4', 'text': stripped[5:].strip()})
        elif stripped.startswith('- ') or stripped.startswith('* '):
            blocks.append({'type': 'bullet', 'text': stripped[2:].strip()})
        elif re.match(r'^\d+\.\s', stripped):
            blocks.append({'type': 'numbered', 'text': re.sub(r'^\d+\.\s', '', stripped)})
        elif stripped.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            blocks.append({'type': 'bold_line', 'text': stripped.strip('*')})
        elif stripped.startswith('---'):
            blocks.append({'type': 'hr'})
        else:
            blocks.append({'type': 'paragraph', 'text': stripped})
        i += 1
    if in_table and table_rows:
        blocks.append({'type': 'table', 'rows': table_rows})
    return blocks


def clean_markdown_inline(text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text


def add_run_with_formatting(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            clean = re.sub(r'\*(.*?)\*', r'\1', part)
            clean = re.sub(r'`(.*?)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean)
            if clean:
                paragraph.add_run(clean)


def generate_docx(md_path, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
        sys.exit(1)

    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_md_to_blocks(md_text)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Local aliases — pinned to brand module (single source of truth)
    DARK_BLUE  = brand.MIDNIGHT_BLUE
    ACCENT     = brand.SKY_BLUE
    MID_GRAY   = brand.GRAY_500
    BLACK      = brand.BLACK

    def sp(para, before=0, after=4):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = DARK_BLUE
        sp(p, before=12, after=8)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = ACCENT
        sp(p, before=14, after=3)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), brand.HEX_SKY_BLUE)
        pBdr.append(bot); pPr.append(pBdr)

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = MID_GRAY
        sp(p, before=8, after=2)

    def add_h4(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = BLACK
        sp(p, before=4, after=2)

    def add_body(text):
        p = doc.add_paragraph()
        add_run_with_formatting(p, text)
        for r in p.runs:
            r.font.size = Pt(10); r.font.color.rgb = BLACK
        sp(p, before=1, after=3)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        add_run_with_formatting(p, text)
        for r in p.runs:
            r.font.size = Pt(10); r.font.color.rgb = BLACK
        sp(p, before=0, after=2)

    def shade_cell(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_table_block(rows):
        if not rows: return
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= num_cols: break
                cell = row.cells[c_idx]
                clean = clean_markdown_inline(cell_text)
                p = cell.paragraphs[0]
                run = p.add_run(clean)
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shade_cell(cell, brand.HEX_MIDNIGHT_BLUE)
                else:
                    run.font.color.rgb = BLACK
                    shade_cell(cell, 'EBF3FA' if r_idx % 2 == 0 else 'FFFFFF')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # Title page
    title_text = "Análisis Competitivo"
    for block in blocks[:5]:
        if block['type'] == 'h1':
            title_text = clean_markdown_inline(block['text'])
            break

    # Brand logo at top of cover (via brand module)
    brand.apply_cover_logo(doc, width_inches=3.8, space_after_pt=36)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title_text)
    run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = DARK_BLUE
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_m = p_meta.add_run(f"Nexostrat — Uso Interno  ·  {datetime.now().strftime('%d %b %Y')}")
    run_m.font.size = Pt(10); run_m.font.italic = True; run_m.font.color.rgb = MID_GRAY

    doc.add_page_break()

    skip_title = True
    for block in blocks:
        btype = block['type']
        if btype == 'h1':
            if skip_title: skip_title = False; continue
            add_h1(block['text'])
        elif btype == 'h2': add_h2(block['text'])
        elif btype == 'h3': add_h3(block['text'])
        elif btype == 'h4': add_h4(block['text'])
        elif btype == 'bullet': add_bullet(block['text'])
        elif btype == 'numbered':
            p = doc.add_paragraph(style='List Number')
            add_run_with_formatting(p, block['text'])
            for r in p.runs: r.font.size = Pt(10); r.font.color.rgb = BLACK
            sp(p, before=0, after=2)
        elif btype == 'table': add_table_block(block['rows'])
        elif btype == 'bold_line':
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(block['text']))
            run.bold = True; run.font.size = Pt(10); run.font.color.rgb = DARK_BLUE
            sp(p, before=4, after=2)
        elif btype == 'paragraph':
            if block['text']: add_body(block['text'])

    # Brand header + footer (via brand module — skips cover automatically)
    brand.apply_brand_header(doc)
    brand.apply_brand_footer(doc)

    doc.save(output_path)
    print(f"✅ DOCX generado: {output_path}")
    print(f"   Tamaño: {Path(output_path).stat().st_size // 1024} KB")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Uso: python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    if not Path(sys.argv[1]).exists():
        print(f"ERROR: No se encontró: {sys.argv[1]}")
        sys.exit(1)
    md_path = sys.argv[1]
    out_path = sys.argv[2]
    generate_docx(md_path, out_path)

    # Plan 02a Task 9 — Baserow deliverables sync (fail-safe).
    # Canonical path layout: pipeline/clients/<slug>/<station>/runs/<date>/<file>.
    # Any failure (path doesn't match, no client row, network error) is logged
    # and swallowed — never blocks the .docx write that just succeeded.
    try:
        import os as _os
        import sys as _sys
        from pathlib import Path as _Path
        _sys.path.insert(0, "/srv/Nexostrat/skills/shared")
        import baserow as _baserow
        if _os.environ.get("BASEROW_URL") and _os.environ.get("BASEROW_API_TOKEN"):
            _parts = _Path(md_path).parts
            if len(_parts) >= 5:
                _slug = _parts[-5]
                _client = _baserow._find_one("clients", "slug", _slug)
                if _client:
                    _baserow.post_deliverable(
                        client_id=_client["id"],
                        skill="competitor-analyst",
                        file_md=str(md_path),
                        file_docx=str(out_path),
                        file_pdf="",
                    )
    except Exception as _e:
        import sys as _sys
        print(f"[baserow-sync] skipped: {_e}", file=_sys.stderr)
