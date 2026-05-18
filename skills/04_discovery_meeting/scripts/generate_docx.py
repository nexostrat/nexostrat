#!/usr/bin/env python3
"""
generate_docx.py — Converts a discovery-meeting .md guión to a styled .docx

Usage:
    python generate_docx.py <input.md> <output.docx>

Requires: python-docx
    pip install python-docx --break-system-packages
"""

import sys
import re
from pathlib import Path
from datetime import datetime


def parse_md_to_blocks(md_text):
    """Parse markdown into structured blocks for docx rendering."""
    blocks = []
    lines = md_text.split('\n')
    i = 0

    # Skip YAML frontmatter
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1

    in_table = False
    table_rows = []
    in_blockquote = False
    bq_lines = []
    bq_color = 'E3F2FD'
    bq_border = '1565C0'

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Blockquote detection ──────────────────────────────────────────
        if stripped.startswith('> ') or stripped == '>':
            content = stripped[2:] if stripped.startswith('> ') else ''
            if not in_blockquote:
                in_blockquote = True
                bq_lines = []
                # Determine color from first content line
                if '🟢' in content or 'APERTURA' in content or 'CREDENCIAL' in content:
                    bq_color = 'E8F5E9'
                    bq_border = '2E7D32'
                elif '⚠️' in content or 'SENSIBLE' in content or 'ZONA' in content:
                    bq_color = 'FFF8E1'
                    bq_border = 'F57F17'
                elif '⏱️' in content or 'TIEMPO' in content or 'TIMING' in content or 'DISTRIBUCIÓN' in content:
                    bq_color = 'F3E5F5'
                    bq_border = '6A1B9A'
                elif 'ℹ️' in content or 'CÓMO USAR' in content or 'Cómo usar' in content:
                    bq_color = 'E3F2FD'
                    bq_border = '1565C0'
                else:
                    bq_color = 'F5F5F5'
                    bq_border = '757575'
            if content:
                bq_lines.append(content)
            i += 1
            continue
        else:
            if in_blockquote and bq_lines:
                blocks.append({'type': 'blockquote', 'lines': bq_lines,
                                'color': bq_color, 'border': bq_border})
                bq_lines = []
                in_blockquote = False

        # ── Table detection ───────────────────────────────────────────────
        if '|' in stripped and stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^\|[-| :]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_rows:
                blocks.append({'type': 'table', 'rows': table_rows})
                table_rows = []
                in_table = False

        # ── Standard blocks ───────────────────────────────────────────────
        if not stripped:
            blocks.append({'type': 'empty'})
        elif stripped.startswith('# ') and not stripped.startswith('##'):
            blocks.append({'type': 'h1', 'text': stripped[2:].strip()})
        elif stripped.startswith('## '):
            blocks.append({'type': 'h2', 'text': stripped[3:].strip()})
        elif stripped.startswith('### '):
            blocks.append({'type': 'h3', 'text': stripped[4:].strip()})
        elif stripped.startswith('#### '):
            blocks.append({'type': 'h4', 'text': stripped[5:].strip()})
        elif stripped.startswith('- ') or stripped.startswith('* '):
            blocks.append({'type': 'bullet', 'text': stripped[2:].strip()})
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            blocks.append({'type': 'numbered', 'text': text})
        elif stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            text = re.sub(r'^- \[.?\] ?', '', stripped)
            blocks.append({'type': 'checklist', 'text': text})
        elif stripped.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
        elif stripped.startswith('---') and len(stripped) >= 3 and set(stripped) == {'-'}:
            blocks.append({'type': 'hr'})
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            blocks.append({'type': 'bold_line', 'text': stripped.strip('*')})
        elif stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            blocks.append({'type': 'italic_line', 'text': stripped.strip('*')})
        else:
            if stripped:
                blocks.append({'type': 'paragraph', 'text': stripped})

        i += 1

    # Flush remaining
    if in_blockquote and bq_lines:
        blocks.append({'type': 'blockquote', 'lines': bq_lines,
                        'color': bq_color, 'border': bq_border})
    if in_table and table_rows:
        blocks.append({'type': 'table', 'rows': table_rows})

    return blocks


def clean_markdown_inline(text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text


def add_run_with_formatting(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            clean = re.sub(r'\*(.*?)\*', r'\1', part)
            clean = re.sub(r'`(.*?)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean)
            if clean:
                paragraph.add_run(clean)


def generate_docx(md_path, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
        sys.exit(1)

    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_md_to_blocks(md_text)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ── Brand colors ──────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x1A, 0x2E, 0x4A)
    ACCENT    = RGBColor(0x00, 0x7A, 0xC3)
    MID_GRAY  = RGBColor(0x55, 0x65, 0x77)
    RED_ALERT = RGBColor(0xC6, 0x28, 0x28)
    BLACK     = RGBColor(0x1A, 0x1A, 0x1A)

    def set_para_spacing(para, before=0, after=4):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        set_para_spacing(p, before=12, after=8)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        set_para_spacing(p, before=14, after=3)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '007AC3')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        # Detect priority markers for area headings
        clean = clean_markdown_inline(text)
        run = p.add_run(clean)
        run.font.size = Pt(11)
        run.font.bold = True
        if '●●●' in text:
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # dark green
        elif '●●' in text:
            run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)  # orange
        elif 'ZONA SENSIBLE' in text or '⚠️' in text:
            run.font.color.rgb = RED_ALERT
        else:
            run.font.color.rgb = MID_GRAY
        set_para_spacing(p, before=8, after=2)
        return p

    def add_h4(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BLACK
        set_para_spacing(p, before=4, after=2)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=1, after=3)
        return p

    def add_italic_body(text):
        p = doc.add_paragraph()
        run = p.add_run(clean_markdown_inline(text))
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = MID_GRAY
        set_para_spacing(p, before=1, after=3)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=0, after=2)
        return p

    def add_numbered(text):
        p = doc.add_paragraph(style='List Number')
        add_run_with_formatting(p, text)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=0, after=2)
        return p

    def add_checklist(text):
        p = doc.add_paragraph()
        run_box = p.add_run('☐  ')
        run_box.font.size = Pt(10)
        run_box.font.color.rgb = MID_GRAY
        add_run_with_formatting(p, text)
        for run in p.runs[1:]:
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        set_para_spacing(p, before=0, after=2)
        return p

    def shade_cell(cell, hex_color='F4F7FB'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_left_border(cell, hex_color, size='18'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'thick')
        left.set(qn('w:sz'), size)
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), hex_color)
        tcBorders.append(left)
        tcPr.append(tcBorders)

    def add_blockquote_box(lines, bg_color, border_color):
        """Render blockquote as a colored single-cell box with thick left border."""
        table = doc.add_table(1, 1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        shade_cell(cell, bg_color)
        set_cell_left_border(cell, border_color)

        first = True
        for line in lines:
            if not line.strip():
                continue
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            add_run_with_formatting(p, line)
            for run in p.runs:
                run.font.size = Pt(9.5)
                if run.bold:
                    run.font.color.rgb = DARK_BLUE
                else:
                    run.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_table_block(rows):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx in range(num_cols):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
                cell = row.cells[c_idx]
                clean = clean_markdown_inline(cell_text)
                p = cell.paragraphs[0]
                run = p.add_run(clean)
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shade_cell(cell, '1A2E4A')
                else:
                    run.font.color.rgb = BLACK
                    shade_cell(cell, 'EBF3FA' if r_idx % 2 == 0 else 'FFFFFF')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Title page ────────────────────────────────────────────────────────
    title_text = "Guión de Reunión de Descubrimiento"
    subtitle_text = ""
    for block in blocks[:6]:
        if block['type'] == 'h1':
            title_text = clean_markdown_inline(block['text'])
            break
    for block in blocks[:8]:
        t = block.get('text', '')
        if 'CONFIDENCIAL' in t or 'Ciudad' in t or ('·' in t and 'Ricardo' in t):
            subtitle_text = clean_markdown_inline(t)
            break

    # Cover background box
    cover_table = doc.add_table(1, 1)
    cover_table.style = 'Table Grid'
    cover_cell = cover_table.rows[0].cells[0]
    shade_cell(cover_cell, '1A2E4A')
    set_cell_left_border(cover_cell, '007AC3', size='24')

    p_conf = cover_cell.paragraphs[0]
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conf = p_conf.add_run('🔒  CONFIDENCIAL — USO INTERNO')
    run_conf.font.size = Pt(9)
    run_conf.font.bold = True
    run_conf.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_conf.paragraph_format.space_before = Pt(8)
    p_conf.paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(title_text)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)

    if subtitle_text:
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(subtitle_text)
        run_sub.font.size = Pt(11)
        run_sub.font.color.rgb = MID_GRAY
        p_sub.paragraph_format.space_after = Pt(4)

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(
        f"Nexostrat — Solo para uso de Ricardo  ·  {datetime.now().strftime('%d %b %Y')}")
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = MID_GRAY

    doc.add_page_break()

    # ── Body ──────────────────────────────────────────────────────────────
    skip_h1 = True
    skip_meta_line = True

    for block in blocks:
        btype = block['type']

        if btype == 'h1':
            if skip_h1:
                skip_h1 = False
                continue
            add_h1(block['text'])

        elif btype == 'h2':
            add_h2(block['text'])

        elif btype == 'h3':
            add_h3(block['text'])

        elif btype == 'h4':
            add_h4(block['text'])

        elif btype == 'bullet':
            add_bullet(block['text'])

        elif btype == 'numbered':
            add_numbered(block['text'])

        elif btype == 'checklist':
            add_checklist(block['text'])

        elif btype == 'table':
            add_table_block(block['rows'])

        elif btype == 'blockquote':
            add_blockquote_box(block['lines'], block['color'], block['border'])

        elif btype == 'bold_line':
            t = block['text']
            # Skip the CONFIDENCIAL subtitle line (already on cover)
            if skip_meta_line and ('CONFIDENCIAL' in t or 'Nexostrat' in t):
                skip_meta_line = False
                continue
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(t))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
            set_para_spacing(p, before=4, after=2)

        elif btype == 'italic_line':
            add_italic_body(block['text'])

        elif btype == 'hr':
            p = doc.add_paragraph()
            set_para_spacing(p, before=2, after=2)

        elif btype == 'paragraph':
            if block['text']:
                # Detect question lines (start with italic marker)
                if block['text'].startswith('*"') or block['text'].startswith('*\''):
                    add_italic_body(block['text'])
                else:
                    add_body(block['text'])

        elif btype == 'empty':
            pass

    # ── Footer ────────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(
            f"GUIÓN CONFIDENCIAL — Nexostrat — Solo para uso de Ricardo — {datetime.now().strftime('%B %Y')}")
        run.font.size = Pt(8)
        run.font.color.rgb = MID_GRAY

    doc.save(output_path)
    print(f"✅ DOCX generado: {output_path}")
    size_kb = Path(output_path).stat().st_size // 1024
    print(f"   Tamaño: {size_kb} KB")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Uso: python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = sys.argv[1]
    out_path = sys.argv[2]

    if not Path(md_path).exists():
        print(f"ERROR: No se encontró el archivo: {md_path}")
        sys.exit(1)

    generate_docx(md_path, out_path)
