#!/usr/bin/env python3
"""
generate_docx.py — Reporte de Oportunidades de IA
Nexostrat

Usage: python generate_docx.py <input.md> <output.docx>
"""

import sys
import re
import json
from datetime import datetime
from pathlib import Path

# Brand surface (palette, logos, header/footer helpers)
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "shared"))
try:
    import brand
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Brand surface — single source of truth in skills/shared/brand.py ──────────
DARK_BLUE   = brand.MIDNIGHT_BLUE
ACCENT      = brand.SKY_BLUE
MID_GRAY    = brand.GRAY_500
LIGHT_GRAY  = RGBColor(0xF4, 0xF5, 0xF7)         # Local: alt table-row tint
WHITE       = brand.WHITE
GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)         # Local: Quick Win indicator
AMBER_DARK  = RGBColor(0xE6, 0x5C, 0x00)         # Local: effort indicator
SUCCESS_GRN = RGBColor(0x2E, 0x7D, 0x32)         # Local: badge green
BRAND_FONT  = brand.BRAND_FONT

# ─── Quadrant colors for 2x2 matrix ────────────────────────────────────────────
Q_VICTORIA   = "D5F5E3"  # top-left: high impact, easy  → green
Q_ESTRATEGIC = "D6EAF8"  # top-right: high impact, hard → blue
Q_RAPIDO     = "FDFEFE"  # bottom-left: low impact, easy → white/light
Q_DESCARTAR  = "FDFEFE"  # bottom-right: low impact, hard → white/light


# ─── XML helpers ────────────────────────────────────────────────────────────────
def hex_to_rgb(hex_str):
    h = hex_str.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, size='12'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')

    def make_border(side, color, val='single'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        return b

    def make_none(side):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        return b

    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            tcBorders.append(make_border(side, color))
        else:
            tcBorders.append(make_none(side))
    # Always suppress inside borders
    tcBorders.append(make_none('insideH'))
    tcBorders.append(make_none('insideV'))
    tcPr.append(tcBorders)

def remove_all_borders(cell):
    set_cell_borders(cell)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    twips = int(width_cm * 567)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cell_text(cell, text, font_size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=0):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def get_circled_number(n):
    """Return Unicode circled digit. ①-⑳ for 1-20, fallback to (n) beyond."""
    if 1 <= n <= 20:
        return chr(0x2460 + n - 1)
    return f'({n})'

def dots_score(val, max_val=5):
    """●●●○○ style display"""
    return '●' * val + '○' * (max_val - val)


# ─── Page setup ─────────────────────────────────────────────────────────────────
def setup_page(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)


# ─── Header / Footer ────────────────────────────────────────────────────────────
def add_header_footer(doc, company_name):
    """Brand header (with company name) + footer (with Confidencial marker)."""
    brand.apply_brand_header(doc, extra=company_name)
    brand.apply_brand_footer(doc, extra="Confidencial")


# ─── Cover page ─────────────────────────────────────────────────────────────────
def add_cover_page(doc, company_name, date_str):
    # Top brand band — Midnight Blue with the full Nexostrat logo overlay
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade_cell(cell, brand.HEX_MIDNIGHT_BLUE)
    set_row_height(t.rows[0], 2.0)
    remove_all_borders(cell)
    # Logo overlay (Midnight-bg transparent — blends with the band)
    brand.insert_logo_in_cell(cell, brand.LOGO_MIDNIGHT, width_inches=3.0)

    # Spacer
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    # Subtitle line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('REPORTE DE OPORTUNIDADES DE IA')
    run.font.size  = Pt(13)
    run.font.color.rgb = ACCENT
    run.font.bold  = True
    run.font.name  = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.font.size  = Pt(28)
    run.font.color.rgb = DARK_BLUE
    run.font.bold  = True
    run.font.name  = 'Calibri'
    p.paragraph_format.space_after = Pt(16)

    # Divider line (using a 1-row table with only bottom border)
    t2 = doc.add_table(1, 1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t2.rows[0].cells[0]
    set_row_height(t2.rows[0], 0.05)
    shade_cell(c, brand.HEX_SKY_BLUE)
    remove_all_borders(c)

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Preparado por Nexostrat')
    run.font.size  = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.font.name  = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str)
    run.font.size  = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.font.name  = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

    # Confidential note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Documento confidencial — Uso exclusivo del cliente')
    run.font.size  = Pt(9)
    run.font.color.rgb = MID_GRAY
    run.font.italic = True
    run.font.name  = 'Calibri'

    doc.add_page_break()


# ─── Section heading ────────────────────────────────────────────────────────────
def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE
    elif level == 2:
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = ACCENT
    else:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = DARK_BLUE


# ─── Body paragraph ─────────────────────────────────────────────────────────────
def add_body(doc, text, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name   = 'Calibri'
    run.font.size   = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


# ─── Callout box ────────────────────────────────────────────────────────────────
def add_callout_box(doc, lines, bg_hex, border_hex):
    t = doc.add_table(1, 1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shade_cell(cell, bg_hex)
    # Thick left border only
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'right', 'bottom', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'thick')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), border_hex)
    tcBorders.append(left)
    tcPr.append(tcBorders)

    cell.text = ''
    for i, line in enumerate(lines):
        p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        # Detect bold **text**
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Calibri'
            else:
                r = p.add_run(part)
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(4)


# ─── Opportunity table ──────────────────────────────────────────────────────────
def add_opportunity_table(doc, opportunities):
    """Styled table for the opportunity inventory."""
    headers = ['ID', 'Oportunidad', 'Área', 'Descripción', 'Impacto', 'Complejidad', 'Riesgo', 'QW Score']
    col_widths = [0.8, 3.0, 2.5, 6.5, 1.5, 1.8, 1.3, 1.5]

    t = doc.add_table(len(opportunities) + 1, len(headers))
    t.style = 'Table Grid'

    # Header row
    for c, (h, w) in enumerate(zip(headers, col_widths)):
        cell = t.rows[0].cells[c]
        shade_cell(cell, brand.HEX_MIDNIGHT_BLUE)
        set_cell_width(cell, w)
        cell_text(cell, h, font_size=9, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)

    # Data rows
    for r, opp in enumerate(opportunities, 1):
        row = t.rows[r]
        # Alternating shading
        bg = 'F0F4F8' if r % 2 == 0 else 'FFFFFF'
        qw = opp['impacto'] + (6 - opp['complejidad']) - opp['riesgo']

        vals = [
            str(opp['id']),
            opp['nombre'],
            opp['area'],
            opp['descripcion'],
            dots_score(opp['impacto']),
            dots_score(opp['complejidad']),
            dots_score(opp['riesgo']),
            str(qw),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for c, (val, w, align) in enumerate(zip(vals, col_widths, aligns)):
            cell = row.cells[c]
            shade_cell(cell, bg)
            set_cell_width(cell, w)
            cell_text(cell, val, font_size=9, align=align, space_before=2, space_after=2)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


# ─── 5×5 Grid chart ─────────────────────────────────────────────────────────────
def add_5x5_grid(doc, opportunities, y_key, title, x_label='Impacto', y_label=None):
    """
    Draws a 5×5 grid chart as a python-docx table.
    Only the left border of col-0 (data) and bottom border of row-4 (data) are black.
    Opportunities shown as circled numbers inside cells.

    Table layout — 6 rows × 6 cols:
      Col 0 (width 0.7cm): Y-axis digit labels (rows 0-4: 5,4,3,2,1; row 5: y_label)
      Cols 1-5 (1.4cm ea): rows 0-4 = grid cells; row 5 = X-axis digits 1-5
    """
    if y_label is None:
        y_label = y_key.capitalize()

    # Build grid content: grid[row][col], row 0 = y=5, row 4 = y=1
    grid = [["" for _ in range(5)] for _ in range(5)]
    for opp in opportunities:
        x = max(1, min(5, opp['impacto']))
        y = max(1, min(5, opp[y_key]))
        col = x - 1
        row = 5 - y
        sym = get_circled_number(opp['id'])
        grid[row][col] = (grid[row][col] + ' ' + sym).strip() if grid[row][col] else sym

    # Chart title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

    # Axis labels note above
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    run2 = p2.add_run(f'Eje X: {x_label} (1=bajo, 5=alto)   |   Eje Y: {y_label} (1=baja, 5=alta)')
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = MID_GRAY
    run2.font.italic = True
    run2.font.name = 'Calibri'

    # Create 6×6 table
    tbl = doc.add_table(6, 6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    Y_VALS = [5, 4, 3, 2, 1]
    CELL_W = 1.4   # data cells
    LABEL_W = 0.8  # label column

    for r in range(6):
        for c in range(6):
            cell = tbl.rows[r].cells[c]
            cell.text = ''
            remove_all_borders(cell)

            if r < 5 and c == 0:
                # Y-axis digit label
                set_cell_width(cell, LABEL_W)
                set_row_height(tbl.rows[r], CELL_W)
                cell_text(cell, str(Y_VALS[r]), font_size=9,
                          color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.RIGHT,
                          space_before=4, space_after=0)

            elif r < 5 and c >= 1:
                # Data cell
                set_cell_width(cell, CELL_W)
                shade_cell(cell, 'F8FAFC')
                content = grid[r][c - 1]
                if content:
                    cell_text(cell, content, font_size=11, bold=True,
                              color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER,
                              space_before=3, space_after=0)

                # L-shaped border logic:
                # Left border on col 1 (first data col) for rows 0-4
                # Bottom border on row 4 (last data row) for cols 1-5
                left_color  = brand.HEX_MIDNIGHT_BLUE if c == 1 else None
                bottom_color = brand.HEX_MIDNIGHT_BLUE if r == 4 else None
                set_cell_borders(cell, left=left_color, bottom=bottom_color, size='18')

            elif r == 5 and c == 0:
                # Corner cell — Y-axis name rotated not possible; leave blank
                set_cell_width(cell, LABEL_W)
                set_row_height(tbl.rows[r], 0.65)

            elif r == 5 and c >= 1:
                # X-axis digit labels
                set_cell_width(cell, CELL_W)
                set_row_height(tbl.rows[r], 0.65)
                cell_text(cell, str(c), font_size=9,
                          color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=2, space_after=0)

    # Axis name labels below
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(4)
    run3 = p3.add_run(f'          ← {y_label}  (eje Y, de abajo hacia arriba)     {x_label} → (eje X, de izquierda a derecha)')
    run3.font.size = Pt(7.5)
    run3.font.color.rgb = MID_GRAY
    run3.font.italic = True
    run3.font.name = 'Calibri'

    # Legend
    legend_p = doc.add_paragraph()
    legend_p.paragraph_format.space_before = Pt(4)
    legend_p.paragraph_format.space_after  = Pt(6)
    for opp in opportunities:
        qw = opp['impacto'] + (6 - opp['complejidad']) - opp['riesgo']
        sym = get_circled_number(opp['id'])
        run = legend_p.add_run(f"{sym} {opp['nombre']}  ")
        run.font.size = Pt(8)
        run.font.color.rgb = MID_GRAY
        run.font.name = 'Calibri'


# ─── 2×2 Priority Matrix ─────────────────────────────────────────────────────────
def add_2x2_matrix(doc, opportunities):
    """
    X axis: Impacto (Bajo 1-3 / Alto 4-5)
    Y axis: Facilidad = 6 - Complejidad (Baja 1-2 / Alta 3-5)

    Quadrants:
      Top-left:     Alto impacto + Alta facilidad  → Victorias Rápidas  (green)
      Top-right:    Bajo impacto + Alta facilidad  → Proyectos de Relleno (light)
      Bottom-left:  Alto impacto + Baja facilidad  → Apuestas Estratégicas (blue)
      Bottom-right: Bajo impacto + Baja facilidad  → Descartar (light gray)
    """
    # Build quadrant assignments
    tl, tr, bl, br = [], [], [], []
    for opp in opportunities:
        high_impact   = opp['impacto'] >= 4
        high_facility = (6 - opp['complejidad']) >= 3
        sym = get_circled_number(opp['id'])
        label = f"{sym} {opp['nombre']}"
        if high_impact and high_facility:
            tl.append(label)
        elif not high_impact and high_facility:
            tr.append(label)
        elif high_impact and not high_facility:
            bl.append(label)
        else:
            br.append(label)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('Gráfica 3: Matriz de Priorización — Impacto vs. Facilidad de Implementación')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

    CELL_W = 7.0

    def quadrant_content(cell, title, items, bg_hex, title_color):
        shade_cell(cell, bg_hex)
        remove_all_borders(cell)
        set_cell_width(cell, CELL_W)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = title_color
        r.font.name = 'Calibri'
        for item in items:
            pi = cell.add_paragraph()
            pi.paragraph_format.space_before = Pt(1)
            pi.paragraph_format.space_after  = Pt(1)
            ri = pi.add_run(f'  {item}')
            ri.font.size = Pt(9)
            ri.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            ri.font.name = 'Calibri'
        if not items:
            pi = cell.add_paragraph()
            ri = pi.add_run('  (ninguna en este cuadrante)')
            ri.font.size = Pt(8.5)
            ri.font.color.rgb = MID_GRAY
            ri.font.italic = True
            ri.font.name = 'Calibri'

    # Row labels
    tbl = doc.add_table(3, 3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Top-left corner: empty (axis labels)
    corner = tbl.rows[0].cells[0]
    corner.text = ''
    remove_all_borders(corner)
    set_cell_width(corner, 1.5)
    set_row_height(tbl.rows[0], 5.5)

    # Top header: "Alta Facilidad"
    th = tbl.rows[0].cells[1]
    th.merge(tbl.rows[0].cells[2])
    th.text = ''
    remove_all_borders(th)
    p = th.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ALTA FACILIDAD DE IMPLEMENTACIÓN')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = SUCCESS_GRN; r.font.name = 'Calibri'

    # Data rows (row 1 = high facility, row 2 = low facility)
    for row_idx, (row_label, items_left, items_right, bg_l, bg_r, tc_l, tc_r) in enumerate([
        ('Alta\nFacilidad', tl, tr, 'D5F5E3', 'FDFEFE', SUCCESS_GRN, MID_GRAY),
        ('Baja\nFacilidad', bl, br, 'D6EAF8', 'FDFEFE', ACCENT, MID_GRAY),
    ], 1):
        # Y-axis label cell
        yl = tbl.rows[row_idx].cells[0]
        yl.text = ''
        remove_all_borders(yl)
        set_cell_width(yl, 1.5)
        p = yl.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(row_label)
        r.font.size = Pt(8.5); r.font.color.rgb = MID_GRAY; r.bold = True; r.font.name = 'Calibri'

        # Left quadrant: high or low impact, current facility row
        titles = {
            (1, True):  '🏆 Victorias Rápidas\n(Alto impacto, fácil de implementar)',
            (1, False): '💤 Proyectos de Relleno\n(Bajo impacto, fácil de implementar)',
            (2, True):  '🎯 Apuestas Estratégicas\n(Alto impacto, requiere más esfuerzo)',
            (2, False): '⏸ Descartar por Ahora\n(Bajo impacto, alta dificultad)',
        }
        quadrant_content(tbl.rows[row_idx].cells[1],
                         titles[(row_idx, True)], items_left, bg_l, tc_l)
        quadrant_content(tbl.rows[row_idx].cells[2],
                         titles[(row_idx, False)], items_right, bg_r, tc_r)

    # X-axis label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('← BAJO IMPACTO          ALTO IMPACTO →')
    r.font.size = Pt(9); r.font.color.rgb = MID_GRAY; r.bold = True; r.font.name = 'Calibri'


# ─── Generic table renderer ─────────────────────────────────────────────────────
def add_generic_table(doc, headers, rows):
    n_cols = len(headers)
    t = doc.add_table(len(rows) + 1, n_cols)
    t.style = 'Table Grid'
    for c, h in enumerate(headers):
        cell = t.rows[0].cells[c]
        shade_cell(cell, brand.HEX_MIDNIGHT_BLUE)
        cell_text(cell, h, font_size=9, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    for r, row_data in enumerate(rows, 1):
        bg = 'F0F4F8' if r % 2 == 0 else 'FFFFFF'
        for c, val in enumerate(row_data):
            cell = t.rows[r].cells[c]
            shade_cell(cell, bg)
            cell_text(cell, str(val), font_size=9, space_before=2, space_after=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Markdown parser ────────────────────────────────────────────────────────────
def parse_callout_type(first_line):
    fl = first_line.upper()
    if '🟢' in first_line or 'INSIGHT' in fl:
        return 'E8F5E9', '2E7D32'
    if '⚠️' in first_line or 'ALERTA' in fl:
        return 'FFF8E1', 'F57F17'
    if '💡' in first_line or 'QUICK WIN' in fl:
        return 'E8F5E9', '1B5E20'
    if 'ℹ️' in first_line or 'NOTA' in fl:
        return 'E3F2FD', '1565C0'
    return 'F3F4F6', brand.HEX_GRAY_500  # Neutral callout (gray border)


def render_inline(p, text):
    """Render a paragraph with **bold** and *italic* inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            r = p.add_run(part)
        r.font.name  = 'Calibri'
        r.font.size  = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def parse_opportunity_table(table_lines):
    """Parse | ID | Oportunidad | Área | Descripción | Impacto | Complejidad | Riesgo | rows."""
    opportunities = []
    for line in table_lines:
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if len(cells) < 7:
            continue
        try:
            opp_id = int(cells[0])
        except ValueError:
            continue
        try:
            impacto     = int(cells[4])
            complejidad = int(cells[5])
            riesgo      = int(cells[6])
        except (ValueError, IndexError):
            continue
        opportunities.append({
            'id':          opp_id,
            'nombre':      cells[1],
            'area':        cells[2],
            'descripcion': cells[3],
            'impacto':     impacto,
            'complejidad': complejidad,
            'riesgo':      riesgo,
        })
    return opportunities


def parse_generic_table(lines):
    """Parse a markdown table into (headers, rows)."""
    headers, rows = [], []
    for line in lines:
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if re.match(r'^[-: ]+$', cells[0]):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def process_markdown(doc, md_text, opportunities_ref):
    """
    Walk through markdown lines and render into the docx.
    opportunities_ref is a list that will be populated when TABLA_OPORTUNIDADES is found.
    Special triggers:
      - ### TABLA_OPORTUNIDADES  → parse table, render it, then insert both 5×5 charts
      - ## 4. Matriz             → after this heading, insert 2×2 matrix
      - > emoji text             → callout box
    """
    lines = md_text.split('\n')
    i = 0
    in_blockquote = False
    bq_lines = []
    bq_bg = bq_border = None
    in_opp_table = False
    opp_table_lines = []
    in_gen_table = False
    gen_table_lines = []
    section4_pending_matrix = False
    skip_cover = True   # skip until first ## 1.

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ── Skip YAML front matter ──
        if stripped == '---' and i < 3:
            i += 1
            continue

        # ── End blockquote ──
        if in_blockquote and not stripped.startswith('>'):
            in_blockquote = False
            add_callout_box(doc, bq_lines, bq_bg, bq_border)
            bq_lines = []

        # ── End opportunity table ──
        if in_opp_table:
            if stripped.startswith('|'):
                opp_table_lines.append(stripped)
                i += 1
                continue
            else:
                # Table ended — parse and render
                opps = parse_opportunity_table(opp_table_lines)
                opportunities_ref.extend(opps)
                add_opportunity_table(doc, opps)
                # Insert both 5×5 charts
                if opps:
                    add_5x5_grid(doc, opps, 'complejidad',
                                 'Gráfica 1: Impacto vs. Complejidad de Implementación',
                                 x_label='Impacto', y_label='Complejidad')
                    add_5x5_grid(doc, opps, 'riesgo',
                                 'Gráfica 2: Impacto vs. Riesgo',
                                 x_label='Impacto', y_label='Riesgo')
                in_opp_table = False
                opp_table_lines = []

        # ── End generic table ──
        if in_gen_table:
            if stripped.startswith('|'):
                gen_table_lines.append(stripped)
                i += 1
                continue
            else:
                headers, rows = parse_generic_table(gen_table_lines)
                if headers:
                    add_generic_table(doc, headers, rows)
                in_gen_table = False
                gen_table_lines = []

        # ── Blockquote ──
        if stripped.startswith('>'):
            content = stripped[2:] if stripped.startswith('> ') else stripped[1:]
            if not in_blockquote:
                in_blockquote = True
                bq_bg, bq_border = parse_callout_type(content)
                bq_lines = []
            if content:
                bq_lines.append(content)
            i += 1
            continue

        # ── Headings ──
        if stripped.startswith('# ') and not stripped.startswith('## '):
            # H1 — document/company title — skip (handled in cover)
            i += 1
            continue

        if stripped.startswith('### TABLA_OPORTUNIDADES'):
            in_opp_table = True
            opp_table_lines = []
            i += 1
            continue

        if stripped.startswith('## '):
            title = stripped.lstrip('# ').strip()
            add_section_heading(doc, title, level=1)
            # If section 4, prepare to insert matrix after any following narrative
            section4_pending_matrix = title.startswith('4.')
            i += 1
            continue

        if stripped.startswith('### '):
            title = stripped.lstrip('# ').strip()
            add_section_heading(doc, title, level=2)
            i += 1
            continue

        if stripped.startswith('#### '):
            title = stripped.lstrip('# ').strip()
            add_section_heading(doc, title, level=3)
            i += 1
            continue

        # ── Matrix trigger: insert 2×2 after "## 4." heading ──
        if section4_pending_matrix and stripped and not stripped.startswith('#'):
            if opportunities_ref:
                add_2x2_matrix(doc, opportunities_ref)
            section4_pending_matrix = False

        # ── Horizontal rule ──
        if stripped.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run('─' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = LIGHT_GRAY
            i += 1
            continue

        # ── Generic table ──
        if stripped.startswith('|') and not in_opp_table:
            in_gen_table = True
            gen_table_lines = [stripped]
            i += 1
            continue

        # ── Numbered list ──
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Cm(0.5)
            render_inline(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        # ── Bullet list ──
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            render_inline(p, content)
            i += 1
            continue

        # ── Empty line ──
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        render_inline(p, stripped)
        i += 1

    # Flush any open blockquote
    if in_blockquote and bq_lines:
        add_callout_box(doc, bq_lines, bq_bg, bq_border)
    # Flush any open table
    if in_opp_table and opp_table_lines:
        opps = parse_opportunity_table(opp_table_lines)
        opportunities_ref.extend(opps)
        add_opportunity_table(doc, opps)
        if opps:
            add_5x5_grid(doc, opps, 'complejidad',
                         'Gráfica 1: Impacto vs. Complejidad de Implementación')
            add_5x5_grid(doc, opps, 'riesgo',
                         'Gráfica 2: Impacto vs. Riesgo')
    if in_gen_table and gen_table_lines:
        headers, rows = parse_generic_table(gen_table_lines)
        if headers:
            add_generic_table(doc, headers, rows)


# ─── Extract metadata from markdown ─────────────────────────────────────────────
def extract_meta(md_text):
    """Extract company name and date from the first few lines."""
    company = 'Empresa'
    date_str = datetime.now().strftime('%d de %B de %Y')
    for line in md_text.split('\n')[:10]:
        stripped = line.strip()
        if stripped.startswith('## ') and not stripped.startswith('## 1.'):
            company = stripped.lstrip('# ').strip()
        if 'fecha' in stripped.lower() or re.search(r'\d{4}', stripped):
            m = re.search(r'\d{1,2} de \w+ de \d{4}', stripped)
            if m:
                date_str = m.group(0)
    return company, date_str


# ─── Main ────────────────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_md  = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])

    if not input_md.exists():
        print(f"Error: {input_md} not found")
        sys.exit(1)

    md_text = input_md.read_text(encoding='utf-8')
    company_name, date_str = extract_meta(md_text)

    doc = Document()
    setup_page(doc)
    add_cover_page(doc, company_name, date_str)
    add_header_footer(doc, company_name)

    opportunities = []
    process_markdown(doc, md_text, opportunities)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    print(f"✅ Reporte generado: {output_docx}")
    print(f"   Empresa: {company_name}")
    print(f"   Oportunidades encontradas: {len(opportunities)}")
    if opportunities:
        sorted_opps = sorted(opportunities, key=lambda o: o['impacto'] + (6 - o['complejidad']) - o['riesgo'], reverse=True)
        qw1 = sorted_opps[0]
        print(f"   Quick Win #1: {qw1['nombre']} (score {qw1['impacto'] + (6 - qw1['complejidad']) - qw1['riesgo']})")
        if len(sorted_opps) > 1:
            qw2 = sorted_opps[1]
            print(f"   Quick Win #2: {qw2['nombre']} (score {qw2['impacto'] + (6 - qw2['complejidad']) - qw2['riesgo']})")


if __name__ == '__main__':
    main()
