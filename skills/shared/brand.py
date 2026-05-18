"""
Nexostrat brand module — single source of truth for .docx rendering.

Source: operations/assets/brand/Nexostrat_Brand_Guide.docx (v1.0, May 2026)
Used by: skills/{01..05}_*/scripts/generate_docx.py

Provides Aurora palette constants, logo asset paths, and three helpers:
    apply_cover_logo(doc, ...)       — insert logo at current position (cover top)
    apply_brand_header(doc, ...)     — body-pages header strip (skips cover via
                                       different_first_page_header_footer=True)
    apply_brand_footer(doc, ...)     — body-pages footer with PAGE field

Per-skill custom layout (CONFIDENCIAL boxes, 2×2 matrices, blockquote callouts)
stays in each skill's own file. This module owns only the brand surface.
"""

from pathlib import Path

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BRAND_GUIDE_VERSION = "1.0"
BRAND_FONT = "Calibri"

# ── Aurora palette (RGBColor — for python-docx run.font.color.rgb) ────────
MIDNIGHT_BLUE = RGBColor(0x0C, 0x1A, 0x2E)   # H1 / headers / dark accents
OCEAN_DEEP    = RGBColor(0x0D, 0x4A, 0x6B)   # H3 / structural
SKY_BLUE      = RGBColor(0x0E, 0xA5, 0xE9)   # H2 / accent lines / links
EMERALD       = RGBColor(0x10, 0xB9, 0x81)   # Validation / success
AMBER_GOLD    = RGBColor(0xF5, 0x9E, 0x0B)   # STATS ONLY — never body text
ARCTIC_WHITE  = RGBColor(0xF0, 0xFB, 0xFF)
GRAY_100      = RGBColor(0xF5, 0xF5, 0xF5)
GRAY_500      = RGBColor(0x6B, 0x72, 0x80)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLACK         = RGBColor(0x1A, 0x1A, 0x1A)

# ── Aurora palette (hex strings — for OXML w:color w:fill etc.) ───────────
HEX_MIDNIGHT_BLUE = '0C1A2E'
HEX_OCEAN_DEEP    = '0D4A6B'
HEX_SKY_BLUE      = '0EA5E9'
HEX_EMERALD       = '10B981'
HEX_AMBER_GOLD    = 'F59E0B'
HEX_ARCTIC_WHITE  = 'F0FBFF'
HEX_GRAY_100      = 'F5F5F5'
HEX_GRAY_500      = '6B7280'
HEX_WHITE         = 'FFFFFF'

# ── Logo assets (resolved relative to repo root) ──────────────────────────
_REPO_ROOT = Path(__file__).resolve().parents[2]
_LOGOS_DIR = _REPO_ROOT / "operations/assets/brand/Logos"

LOGO_ARCTIC     = _LOGOS_DIR / "Nexostrat_Logo_Fondo_Arctic_Transparente.png"
LOGO_MIDNIGHT   = _LOGOS_DIR / "Nexostrat_Logo_Fondo_Midnight_Transparente.png"
LOGO_BLANCO     = _LOGOS_DIR / "Nexostrat_Logo_Fondo_Blanco_Transparente.png"
LOGO_SKYBLUE    = _LOGOS_DIR / "Nexostrat_Logo_Fondo_SkyBlue_Transparente.png"
LOGO_MONO_DARK  = _LOGOS_DIR / "Nexostrat_Logo_Monocromatico_Oscuro_Transparente.png"
LOGO_MONO_LIGHT = _LOGOS_DIR / "Nexostrat_Logo_Monocromatico_Claro_Transparente.png"


# ── Helpers ───────────────────────────────────────────────────────────────

def apply_cover_logo(doc, logo_path=None, width_inches=3.8,
                     space_before_pt=36, space_after_pt=24):
    """Insert the brand logo at the current document position (typically cover top).

    Defaults to LOGO_ARCTIC (transparent — works on white pages). For dark-bg
    cover bands (e.g. Skill 05's Midnight Blue band), pass LOGO_MIDNIGHT.
    Silently no-ops if the asset is missing.
    """
    if logo_path is None:
        logo_path = LOGO_ARCTIC
    logo_path = Path(logo_path)
    if not logo_path.exists():
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.add_run().add_picture(str(logo_path), width=Inches(width_inches))
    return p


def insert_logo_in_cell(cell, logo_path, width_inches=3.0,
                        align=WD_ALIGN_PARAGRAPH.CENTER):
    """Insert a logo image into an existing table cell (used for cover bands).

    Clears the cell's default text first so the logo is the only content.
    """
    logo_path = Path(logo_path)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if logo_path.exists():
        p.add_run().add_picture(str(logo_path), width=Inches(width_inches))
    return p


def apply_brand_header(doc, label="NEXOSTRAT", extra=None,
                       align=WD_ALIGN_PARAGRAPH.LEFT):
    """Apply the body-pages header strip.

    Sets section.different_first_page_header_footer = True so the cover page
    is excluded. Renders the label (+ optional `extra` after a center-dot)
    in Midnight Blue Calibri Bold 11pt with a Sky Blue 0.75pt bottom rule.
    """
    text = label if not extra else f"{label}  ·  {extra}"
    for section in doc.sections:
        section.different_first_page_header_footer = True
        hp = section.header.paragraphs[0]
        hp.alignment = align
        # Clear existing runs so this is idempotent
        for run in list(hp.runs):
            run._element.getparent().remove(run._element)
        run = hp.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = MIDNIGHT_BLUE
        run.font.name = BRAND_FONT
        # Sky Blue bottom rule (replace any prior pBdr to stay idempotent)
        pPr = hp._p.get_or_add_pPr()
        existing = pPr.find(qn('w:pBdr'))
        if existing is not None:
            pPr.remove(existing)
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), HEX_SKY_BLUE)
        pBdr.append(bot)
        pPr.append(pBdr)


def apply_brand_footer(doc, extra=None):
    """Apply the body-pages footer.

    Renders `nexostrat.com  ·  [<extra>  ·  ]Pág. <PAGE field>` centered in
    Gray 500 Calibri 9pt. Idempotent (clears prior runs).
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in list(fp.runs):
            run._element.getparent().remove(run._element)
        prefix = "nexostrat.com  ·  "
        if extra:
            prefix += f"{extra}  ·  "
        prefix += "Pág. "
        run_f = fp.add_run(prefix)
        run_f.font.size = Pt(9)
        run_f.font.color.rgb = GRAY_500
        run_f.font.name = BRAND_FONT
        page_run = fp.add_run()
        page_run.font.size = Pt(9)
        page_run.font.color.rgb = GRAY_500
        page_run.font.name = BRAND_FONT
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        page_run._r.append(fld_begin)
        page_run._r.append(instr)
        page_run._r.append(fld_end)
